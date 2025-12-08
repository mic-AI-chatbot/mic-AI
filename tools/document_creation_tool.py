import logging
import os
import json
from typing import List, Dict, Any, Optional, Union
from tools.base_tool import BaseTool

# Conditional imports for external libraries
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    DOCX_AVAILABLE = False
    logging.warning("python-docx library not found. DOCX document creation will not be available.")

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except ImportError:
    letter = None
    canvas = None
    REPORTLAB_AVAILABLE = False
    logging.warning("reportlab library not found. PDF document creation will not be available.")

logger = logging.getLogger(__name__)

class DocumentCreationTool(BaseTool):
    """
    A tool for creating documents in various formats (DOCX, Markdown, HTML, PDF).
    """

    def __init__(self, tool_name: str = "document_creation_tool"):
        super().__init__(tool_name)

    @property
    def description(self) -> str:
        return "Creates documents in DOCX, Markdown, HTML, or PDF formats with specified title and content."

    @property
    def parameters(self) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "format": {
                    "type": "string",
                    "description": "The format of the document to create.",
                    "enum": ["docx", "markdown", "html", "pdf"]
                },
                "file_path": {
                    "type": "string",
                    "description": "The absolute path to save the document (e.g., '/path/to/document.docx')."
                },
                "title": {"type": "string", "description": "The title of the document."},
                "content": {"type": "string", "description": "The main content of the document."}
            },
            "required": ["format", "file_path", "content"]
        }

    def _ensure_dir_exists(self, file_path: str) -> None:
        output_dir = os.path.dirname(file_path)
        if output_dir: os.makedirs(output_dir, exist_ok=True)

    def _create_docx(self, file_path: str, title: Optional[str] = None, content: Optional[str] = None) -> Dict[str, Any]:
        if not DOCX_AVAILABLE: raise ImportError("The 'python-docx' library is required for DOCX creation.")
        if not os.path.isabs(file_path): raise ValueError("File path must be an absolute path.")
        self._ensure_dir_exists(file_path)

        document = Document()
        if title: document.add_heading(title, level=1)
        if content: document.add_paragraph(content)
        document.save(file_path)
        return {"status": "success", "message": f"Document '{file_path}' created successfully."}

    def _create_markdown(self, file_path: str, title: Optional[str] = None, content: Optional[str] = None) -> Dict[str, Any]:
        if not os.path.isabs(file_path): raise ValueError("File path must be an absolute path.")
        self._ensure_dir_exists(file_path)

        with open(file_path, 'w', encoding='utf-8') as f:
            if title: f.write(f"# {title}\n\n")
            if content: f.write(content)
        return {"status": "success", "message": f"Document '{file_path}' created successfully."}

    def _create_html(self, file_path: str, title: Optional[str] = None, content: Optional[str] = None) -> Dict[str, Any]:
        if not os.path.isabs(file_path): raise ValueError("File path must be an absolute path.")
        self._ensure_dir_exists(file_path)

        html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset=\"UTF-8\">
    <title>{title if title else "Document"}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; line-height: 1.6; }}
        h1 {{ color: #333; }}
        p {{ color: #555; }}
    </style>
</head>
<body>
    <h1>{title if title else "Document"}</h1>
    <p>{content if content else ""}</p>
</body>
</html>
"""
        with open(file_path, 'w', encoding='utf-8') as f: f.write(html_content)
        return {"status": "success", "message": f"Document '{file_path}' created successfully."}

    def _create_pdf(self, file_path: str, title: Optional[str] = None, content: Optional[str] = None) -> Dict[str, Any]:
        if not REPORTLAB_AVAILABLE: raise ImportError("The 'reportlab' library is required for PDF creation.")
        if not os.path.isabs(file_path): raise ValueError("File path must be an absolute path.")
        self._ensure_dir_exists(file_path)

        c = canvas.Canvas(file_path, pagesize=letter)
        y_position = 750
        if title: c.setFont("Helvetica-Bold", 24); c.drawString(100, y_position, title); y_position -= 30
        if content:
            c.setFont("Helvetica", 12); textobject = c.beginText(); textobject.setTextOrigin(100, y_position)
            for line in content.split('\n'): textobject.textLine(line)
            c.drawText(textobject)
        c.save()
        return {"status": "success", "message": f"Document '{file_path}' created successfully."}

    def execute(self, format: str, file_path: str, content: str, title: Optional[str] = None, **kwargs: Any) -> Dict[str, Any]:
        if format == "docx":
            return self._create_docx(file_path, title, content)
        elif format == "markdown":
            return self._create_markdown(file_path, title, content)
        elif format == "html":
            return self._create_html(file_path, title, content)
        elif format == "pdf":
            return self._create_pdf(file_path, title, content)
        else:
            raise ValueError(f"Unsupported document format: {format}")

if __name__ == '__main__':
    print("Demonstrating DocumentCreationTool functionality...")
    tool = DocumentCreationTool()
    
    output_dir = os.path.abspath("generated_documents")
    
    try:
        os.makedirs(output_dir, exist_ok=True)

        sample_title = "My Sample Document"
        sample_content = "This is some sample content for the document."

        print("\n--- Creating DOCX Document ---")
        docx_path = os.path.join(output_dir, "sample_document.docx")
        result = tool.execute(format="docx", file_path=docx_path, title=sample_title, content=sample_content)
        print(json.dumps(result, indent=2))

        print("\n--- Creating Markdown Document ---")
        md_path = os.path.join(output_dir, "sample_document.md")
        result = tool.execute(format="markdown", file_path=md_path, title=sample_title, content=sample_content)
        print(json.dumps(result, indent=2))

    except Exception as e:
        print(f"An error occurred: {e}")
    finally:
        if os.path.exists(output_dir): shutil.rmtree(output_dir)
        print("\nCleanup complete.")

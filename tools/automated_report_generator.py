import json
import os
import logging
from typing import Union, List, Dict, Any
from tools.base_tool import BaseTool

# ReportLab for PDF generation
try:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    SimpleDocTemplate = Paragraph = Spacer = Table = TableStyle = getSampleStyleSheet = inch = colors = None
    REPORTLAB_AVAILABLE = False
    logging.warning("ReportLab library not found. PDF report generation will not be available. Please install it with 'pip install reportlab'.")

# python-docx for DOCX generation
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    Document = Inches = None
    DOCX_AVAILABLE = False
    logging.warning("python-docx library not found. DOCX report generation will not be available. Please install it with 'pip install python-docx'.")

logger = logging.getLogger(__name__)

class GeneratePDFReportTool(BaseTool):
    """Generates a PDF report from structured data."""
    def __init__(self, tool_name="generate_pdf_report"):
        super().__init__(tool_name=tool_name)

    @property
    def description(self) -> str:
        return "Generates a PDF report from structured data (string, list of dictionaries, or dictionary) with a given title and saves it to a specified file path. Requires ReportLab."

    @property
    def parameters(self) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "The absolute path where the PDF report will be saved."},
                "title": {"type": "string", "description": "The title of the report."},
                "data": {
                    "type": ["string", "array", "object"],
                    "description": "The data to include in the report, either as a JSON string, a list of dictionaries, or a dictionary."
                }
            },
            "required": ["file_path", "title", "data"]
        }

    def execute(self, file_path: str, title: str, data: Union[str, List[Dict[str, Any]], Dict[str, Any]], **kwargs: Any) -> str:
        if not REPORTLAB_AVAILABLE:
            return json.dumps({"error": "ReportLab library is not installed. PDF report generation is not available."})
        
        if not os.path.isabs(file_path):
            return json.dumps({"error": "'file_path' must be an absolute path."})

        doc = SimpleDocTemplate(file_path)
        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph(title, styles['h1']))
        story.append(Spacer(1, 0.2 * inch))

        # Convert data to appropriate format if it's a JSON string
        if isinstance(data, str):
            try:
                data = json.loads(data)
            except json.JSONDecodeError:
                story.append(Paragraph(data, styles['Normal']))
                data = None # Mark as handled if it was a plain string

        if isinstance(data, dict):
            for key, value in data.items():
                story.append(Paragraph(f"<b>{key}:</b> {value}", styles['Normal']))
        elif isinstance(data, list) and all(isinstance(item, dict) for item in data):
            if data:
                headers = list(data[0].keys())
                table_data = [headers]
                for item in data:
                    table_data.append([str(item.get(h, '')) for h in headers])
                
                table = Table(table_data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
        elif data is not None: # If it was a string but not JSON, it's already added
            story.append(Paragraph(str(data), styles['Normal']))

        try:
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            doc.build(story)
            return json.dumps({"message": f"PDF report '{title}' generated and saved to '{os.path.abspath(file_path)}'.", "file_path": os.path.abspath(file_path)})
        except Exception as e:
            logger.error(f"Error generating PDF report: {e}")
            return json.dumps({"error": f"Error generating PDF report: {e}"})

class GenerateHTMLReportTool(BaseTool):
    """Generates an HTML report from structured data."""
    def __init__(self, tool_name="generate_html_report"):
        super().__init__(tool_name=tool_name)

    @property
    def description(self) -> str:
        return "Generates an HTML report from structured data (string, list of dictionaries, or dictionary) with a given title and saves it to a specified file path."

    @property
    def parameters(self) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "The absolute path where the HTML report will be saved."},
                "title": {"type": "string", "description": "The title of the report."},
                "data": {
                    "type": ["string", "array", "object"],
                    "description": "The data to include in the report, either as a JSON string, a list of dictionaries, or a dictionary."
                }
            },
            "required": ["file_path", "title", "data"]
        }

    def execute(self, file_path: str, title: str, data: Union[str, List[Dict[str, Any]], Dict[str, Any]], **kwargs: Any) -> str:
        if not os.path.isabs(file_path):
            return json.dumps({"error": "'file_path' must be an absolute path."})

        html_content = f"<!DOCTYPE html>\n<html>\n<head><title>{title}</title>"
        html_content += "<style>body{font-family:sans-serif; margin:20px;} h1{color:#333;} table{width:100%; border-collapse:collapse; margin-top:20px;} th,td{border:1px solid #ddd; padding:10px; text-align:left;} th{background-color:#f2f2f2; color:#555;}</style>"
        html_content += "</head>\n<body>\n"
        html_content += f"<h1>{title}</h1>\n"

        # Convert data to appropriate format if it's a JSON string
        if isinstance(data, str):
            try:
                data = json.loads(data)
            except json.JSONDecodeError:
                html_content += f"<p>{data}</p>\n"
                data = None # Mark as handled
        
        if isinstance(data, dict):
            html_content += "<ul>\n"
            for key, value in data.items():
                html_content += f"<li><b>{key}:</b> {value}</li>\n"
            html_content += "</ul>\n"
        elif isinstance(data, list) and all(isinstance(item, dict) for item in data):
            if data:
                headers = list(data[0].keys())
                html_content += "<table><thead><tr>"
                for header in headers:
                    html_content += f"<th>{header}</th>"
                html_content += "</tr></thead><tbody>"
                for item in data:
                    html_content += "<tr>"
                    for header in headers:
                        html_content += f"<td>{item.get(header, '')}</td>"
                    html_content += "</tr>"
                html_content += "</tbody></table>\n"
        elif data is not None: # If it was a string but not JSON, it's already added
            html_content += f"<p>{data}</p>\n"

        html_content += "</body>\n</html>"

        try:
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(html_content)
            return json.dumps({"message": f"HTML report '{title}' generated and saved to '{os.path.abspath(file_path)}'.", "file_path": os.path.abspath(file_path)})
        except Exception as e:
            logger.error(f"Error generating HTML report: {e}")
            return json.dumps({"error": f"Error generating HTML report: {e}"})

class GenerateDOCXReportTool(BaseTool):
    """Generates a DOCX report from structured data."""
    def __init__(self, tool_name="generate_docx_report"):
        super().__init__(tool_name=tool_name)

    @property
    def description(self) -> str:
        return "Generates a DOCX report from structured data (string, list of dictionaries, or dictionary) with a given title and saves it to a specified file path. Requires python-docx."

    @property
    def parameters(self) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "The absolute path where the DOCX report will be saved."},
                "title": {"type": "string", "description": "The title of the report."},
                "data": {
                    "type": ["string", "array", "object"],
                    "description": "The data to include in the report, either as a JSON string, a list of dictionaries, or a dictionary."
                }
            },
            "required": ["file_path", "title", "data"]
        }

    def execute(self, file_path: str, title: str, data: Union[str, List[Dict[str, Any]], Dict[str, Any]], **kwargs: Any) -> str:
        if not DOCX_AVAILABLE:
            return json.dumps({"error": "python-docx library is not installed. DOCX report generation is not available."})
        
        if not os.path.isabs(file_path):
            return json.dumps({"error": "'file_path' must be an absolute path."})

        document = Document()
        document.add_heading(title, level=1)

        # Convert data to appropriate format if it's a JSON string
        if isinstance(data, str):
            try:
                data = json.loads(data)
            except json.JSONDecodeError:
                document.add_paragraph(data)
                data = None # Mark as handled
        
        if isinstance(data, dict):
            for key, value in data.items():
                document.add_paragraph(f"{key}: {value}")
        elif isinstance(data, list) and all(isinstance(item, dict) for item in data):
            if data:
                headers = list(data[0].keys())
                table = document.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                
                for item in data:
                    row_cells = table.add_row().cells
                    for i, header in enumerate(headers):
                        row_cells[i].text = str(item.get(header, ''))
        elif data is not None: # If it was a string but not JSON, it's already added
            document.add_paragraph(str(data))

        try:
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            document.save(file_path)
            return json.dumps({"message": f"DOCX report '{title}' generated and saved to '{os.path.abspath(file_path)}'.", "file_path": os.path.abspath(file_path)})
        except Exception as e:
            logger.error(f"Error generating DOCX report: {e}")
            return json.dumps({"error": f"Error generating DOCX report: {e}"})